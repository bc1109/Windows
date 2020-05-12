import docx


# d = doc.Document('C:\\users\\admin\bcpy\\demo.docx')


# d.paragraphs list of paragraphs


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(getText('C:\\users\\admin\BCPY\\demo.docx'))
